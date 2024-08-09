import numpy as np
import matplotlib.pyplot as plt
from sklearn.metrics import accuracy_score, precision_score, recall_score, f1_score, confusion_matrix, classification_report
from skimage.io import imread
from skimage.transform import resize
from sklearn import svm
import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog
from PIL import Image, ImageTk
from docx import Document
import os
from datetime import datetime

# Load data
X_train = np.load('X_train.npy')
X_test = np.load('X_test.npy')
y_train = np.load('y_train.npy')
y_test = np.load('y_test.npy')
categories = np.load('categories.npy', allow_pickle=True)

# Train the model
model = svm.SVC(kernel='linear', C=1)
model.fit(X_train, y_train)
y_pred = model.predict(X_test)

# Evaluate the model
accuracy = accuracy_score(y_test, y_pred)
precision = precision_score(y_test, y_pred, average='weighted')
recall = recall_score(y_test, y_pred, average='weighted')
f1 = f1_score(y_test, y_pred, average='weighted')

print(f"Accuracy: {accuracy}")
print(f"Precision: {precision}")
print(f"Recall: {recall}")
print(f"F1 Score: {f1}")

# Confusion matrix and classification report
conf_matrix = confusion_matrix(y_test, y_pred)
class_report = classification_report(y_test, y_pred)

print("Confusion Matrix:")
print(conf_matrix)
print("\nClassification Report:")
print(class_report)

plt.figure(figsize=(10, 7))
plt.imshow(conf_matrix, cmap='Blues')
plt.title('Confusion Matrix')
plt.colorbar()
plt.xlabel('Predicted')
plt.ylabel('Actual')
plt.show()

# Define disease descriptions and symptoms
disease_info = {
    'glioma': {
        'description': 'Glioma is a type of tumor that occurs in the brain and spinal cord. Gliomas begin in the glial cells that surround nerve cells and help them function.',
        'symptoms': 'Symptoms include headaches, nausea, vomiting, and seizures.'
    },
    'no_tumor': {
        'description': 'No tumor detected.',
        'symptoms': 'No symptoms related to brain tumors.'
    },
    'meningioma': {
        'description': 'Meningioma is a tumor that arises from the meninges — the membranes that surround the brain and spinal cord. Most meningiomas are noncancerous.',
        'symptoms': 'Symptoms include changes in vision, headaches, and memory loss.'
    },
    'pituitary': {
        'description': 'Pituitary tumors are abnormal growths that develop in the pituitary gland. Some pituitary tumors can cause the gland to produce lower or higher levels of hormones.',
        'symptoms': 'Symptoms include headaches, vision problems, and hormonal imbalances.'
    }
}

# Predict function
def predict_image(model, img_path, categories):
    try:
        img = imread(img_path, as_gray=True)
        img_resized = resize(img, (48, 48)).flatten()
        prediction = model.predict([img_resized])
        predicted_category = categories[prediction[0]]
        return predicted_category
    except Exception as e:
        print(f"Error predicting image: {e}")
        return "Error"

# Browse image function
def browse_image():
    img_path = filedialog.askopenfilename(filetypes=[("Image files", "*.jpg *.jpeg *.png")])
    if img_path:
        predicted_category = predict_image(model, img_path, categories)
        result_label.config(text=f"Predicted class: {predicted_category}")
        save_button.config(state=tk.NORMAL)
        display_image(img_path)

# Display image function
def display_image(img_path):
    try:
        img = Image.open(img_path)
        img = img.resize((250, 250), Image.LANCZOS)
        img_tk = ImageTk.PhotoImage(img)
        panel.configure(image=img_tk)
        panel.image = img_tk
    except Exception as e:
        print(f"Error displaying image: {e}")

# Function to open the input form
def open_input_form():
    input_form = tk.Toplevel(root)
    input_form.title("Patient Information")

    tk.Label(input_form, text="Patient Name:").grid(row=0, column=0)
    tk.Label(input_form, text="Sex:").grid(row=1, column=0)
    tk.Label(input_form, text="Name of Guardian:").grid(row=2, column=0)
    tk.Label(input_form, text="Address:").grid(row=3, column=0)
    tk.Label(input_form, text="Contact No:").grid(row=4, column=0)
    tk.Label(input_form, text="Relationship of Guardian:").grid(row=5, column=0)

    patient_name_entry = tk.Entry(input_form)
    sex_entry = tk.Entry(input_form)
    guardian_name_entry = tk.Entry(input_form)
    address_entry = tk.Entry(input_form)
    contact_no_entry = tk.Entry(input_form)
    relationship_entry = tk.Entry(input_form)

    patient_name_entry.grid(row=0, column=1)
    sex_entry.grid(row=1, column=1)
    guardian_name_entry.grid(row=2, column=1)
    address_entry.grid(row=3, column=1)
    contact_no_entry.grid(row=4, column=1)
    relationship_entry.grid(row=5, column=1)

    def submit_form():
        patient_info = {
            'patient_name': patient_name_entry.get(),
            'sex': sex_entry.get(),
            'guardian_name': guardian_name_entry.get(),
            'address': address_entry.get(),
            'contact_no': contact_no_entry.get(),
            'relationship': relationship_entry.get(),
            'date': datetime.now().strftime("%Y-%m-%d")
        }
        save_prediction(patient_info)
        input_form.destroy()

    submit_button = tk.Button(input_form, text="Submit", command=submit_form)
    submit_button.grid(row=6, column=1)

# Function to save prediction to Word document
def save_prediction(patient_info):
    disease_name = result_label.cget("text").replace("Predicted class: ", "")
    if not disease_name:
        messagebox.showwarning("No Prediction", "No prediction to save.")
        return

    description = disease_info[disease_name]['description']
    symptoms = disease_info[disease_name]['symptoms']

    # Load or create a Word document
    try:
        doc_path = 'prescription.docx'
        
        if os.path.exists(doc_path):
            doc = Document(doc_path)
        else:
            doc = Document()
            doc.add_heading('Patient Prescription', level=1)

        # Add patient information
        doc.add_paragraph(f"Date: {patient_info['date']}")
        doc.add_paragraph(f"Patient Name: {patient_info['patient_name']}")
        doc.add_paragraph(f"Sex: {patient_info['sex']}")
        doc.add_paragraph(f"Name of Guardian: {patient_info['guardian_name']}")
        doc.add_paragraph(f"Address: {patient_info['address']}")
        doc.add_paragraph(f"Contact No: {patient_info['contact_no']}")
        doc.add_paragraph(f"Relationship of Guardian: {patient_info['relationship']}")
        
        # Add or update the disease name, description, and symptoms
        doc.add_paragraph(f"Disease Name: {disease_name}")
        doc.add_paragraph(f"Description: {description}")
        doc.add_paragraph(f"Symptoms: {symptoms}")

        # Save the Word document
        doc.save(doc_path)

        messagebox.showinfo("Success", f"Prediction saved to Word document: {doc_path}")
    except Exception as e:
        messagebox.showerror("Error", f"Failed to save prediction: {e}")

# Set up the GUI
root = tk.Tk()
root.title("Brain Tumor Classification")

frame = tk.Frame(root)
frame.pack(pady=20)

browse_button = tk.Button(frame, text="Browse Image", command=browse_image)
browse_button.pack(side="left", padx=10)

save_button = tk.Button(frame, text="Save Prediction", command=open_input_form, state=tk.DISABLED)
save_button.pack(side="right", padx=10)

result_label = tk.Label(root, text="", font=('Helvetica', 14))
result_label.pack(pady=20)

panel = tk.Label(root)
panel.pack(pady=20)

root.mainloop()
